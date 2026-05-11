import pandas as pd
from docxtpl import DocxTemplate
from docx import Document
import os

def replace_literal_text(doc, old_text, new_text):
    """手動替換文件中所有出現的文字（包含段落、表格、頁首頁尾）"""
    
    def process_container(container):
        # 處理容器中的所有段落
        for p in container.paragraphs:
            if old_text in p.text:
                p.text = p.text.replace(old_text, new_text)
        # 處理容器中的所有表格
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_container(cell)

    # 1. 處理主要內容 (段落與表格)
    process_container(doc)

    # 2. 處理所有節的頁首與頁尾
    for section in doc.sections:
        process_container(section.header)
        process_container(section.footer)
        # 處理首頁不同或奇偶頁不同的情況
        process_container(section.first_page_header)
        process_container(section.first_page_footer)
        process_container(section.even_page_header)
        process_container(section.even_page_footer)


def main():
    excel_file = 'data.xlsx'
    template_file = 'RD260309A09-PSG_HA2026-HC_DVT Test Report Ver A_Template.docx'
    output_file = 'RD260309A09-PSG_HA2026-HC_DVT Test Report Ver A_new_values.docx'

    if not os.path.exists(excel_file):
        print(f"Error: Cannot find {excel_file}")
        return
    if not os.path.exists(template_file):
        print(f"Error: Cannot find {template_file}")
        return

    print(f"Reading {excel_file}...")
    df = pd.read_excel(excel_file)
    
    context = {}
    replacements = []
    for _, row in df.iterrows():
        tag = str(row['Tag_Name']).strip()
        val = str(row['Value']).strip()
        context[tag] = val
        replacements.append((tag, val))

    print(f"Loading template {template_file}...")
    # Method 1: Use docxtpl for {{ Tag_Name }}
    try:
        doc_tpl = DocxTemplate(template_file)
        doc_tpl.render(context)
        doc_tpl.save(output_file)
        print("Completed {{ Tag_Name }} format replacement.")
    except Exception as e:
        print(f"docxtpl rendering failed: {e}")
        # If failed, copy template as output for manual replacement
        import shutil
        shutil.copy(template_file, output_file)

    # Method 2: Handle cases without {{ }} (direct text replacement)
    print("Checking for unbraced text to replace...")
    doc = Document(output_file)
    for tag, val in replacements:
        replace_literal_text(doc, tag, val)
    
    doc.save(output_file)
    print(f"Done! Result saved to: {output_file}")

if __name__ == "__main__":
    main()
