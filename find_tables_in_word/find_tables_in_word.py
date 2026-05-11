import os
import sys
import io
from docx import Document
from docx.text.paragraph import Paragraph
import win32com.client as win32

# 強制設定輸出編碼為 UTF-8，解決 Windows 終端機編碼問題
if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def convert_doc_to_docx(doc_path):
    """將 .doc 檔案轉換為 .docx 檔案，以便 python-docx 讀取"""
    abs_path = os.path.abspath(doc_path)
    docx_path = abs_path + "x"
    
    if os.path.exists(docx_path):
        return docx_path
    
    print(f"正在將 {os.path.basename(doc_path)} 轉換為 .docx...")
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(abs_path)
    # FileFormat=16 為 docx 格式
    doc.SaveAs2(docx_path, FileFormat=16)
    doc.Close()
    # word.Quit() # 這裡不退出，方便後續可能的使用，或者在 main 結束時退出
    return docx_path

def get_outline_level(paragraph):
    """從段落的 XML 中取得大綱層級 (Outline Level)"""
    pPr = paragraph._element.pPr
    if pPr is not None:
        # 定義 XML 命名空間
        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        outlineLvl = pPr.find(f'{ns}outlineLvl')
        if outlineLvl is not None:
            val = outlineLvl.get(f'{ns}val')
            if val is not None:
                return int(val)
    return None

def find_tables_with_headings(file_path):
    """讀取 Word 檔並找出表格及其所屬的大綱標題"""
    if not os.path.exists(file_path):
        print(f"錯誤: 找不到檔案 {file_path}")
        return

    # 如果是 .doc 則嘗試轉換
    if file_path.lower().endswith('.doc'):
        try:
            file_path = convert_doc_to_docx(file_path)
        except Exception as e:
            print(f"轉換失敗: {e}")
            print("請手動將檔案另存為 .docx 後再執行。")
            return

    doc = Document(file_path)
    
    print(f"正在分析文件: {os.path.basename(file_path)}")
    print("=" * 50)
    
    current_heading = "文件開頭 (未分類)"
    table_count = 0
    
    # 遍歷文件 body 中的所有元素 (段落與表格)
    # 使用底層 element.body 遍歷可以維持文件中的先後順序
    for element in doc.element.body:
        # 判斷元素標籤
        tag = element.tag
        
        # 如果是段落 (Paragraph)
        if tag.endswith('p'):
            para = Paragraph(element, doc)
            # 優先檢查 XML 中的大綱層級 (0-8 為標題)
            outline_lvl = get_outline_level(para)
            
            # 檢查樣式名稱
            style_name = para.style.name
            
            # 如果大綱層級存在，或者樣式名稱包含標題關鍵字
            if outline_lvl is not None or style_name.startswith('Heading') or '標題' in style_name:
                text = para.text.strip()
                # 只有當內容不為空時才更新標題
                if text:
                    current_heading = text
                elif outline_lvl is not None:
                    # 如果有層級但沒文字，可能是手動設定層級的空行，暫不更新或記錄
                    pass
        
        # 如果是表格 (Table)
        elif tag.endswith('tbl'):
            table_count += 1
            # 找到該表格在 doc.tables 中的對應物件 (依序排列)
            # 因為我們是順著遍歷，所以這裡可以用 table_count-1 來索引
            table = doc.tables[table_count - 1]
            
            rows = len(table.rows)
            cols = len(table.columns)
            
            print(f"\n[表格 {table_count}]")
            print(f"  位置: 位於標題「{current_heading}」之下")
            print(f"  規格: {rows} 列 x {cols} 欄")
            
            # 預覽第一列內容
            try:
                first_row_cells = [cell.text.strip() for cell in table.rows[0].cells]
                preview = " | ".join(first_row_cells[:3]) # 只預覽前三格
                if len(first_row_cells) > 3:
                    preview += " ..."
                print(f"  內容預覽: {preview}")
            except Exception:
                print("  內容預覽: (讀取失敗或表格為空)")

    print("\n" + "=" * 50)
    print(f"分析完成，共發現 {table_count} 個表格。")

if __name__ == "__main__":
    # 這裡可以修改成你想要讀取的檔案名稱
    target_file = "Template.doc"
    
    # 如果目錄下有 Template.docx 就優先使用
    if os.path.exists("Template.docx"):
        target_file = "Template.docx"
        
    find_tables_with_headings(target_file)
