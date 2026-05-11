import pandas as pd
from docx import Document

def run_based_replace(file_path, excel_path, output_path):
    # 1. 讀取 Excel 內容
    df = pd.read_excel(excel_path)
    replacements = dict(zip(df.iloc[:, 0].astype(str), df.iloc[:, 1].astype(str)))

    # 2. 開啟 Word
    doc = Document(file_path)

    # 定義內部的 Run 替換邏輯
    def replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            # 先檢查整個段落文字裡是否包含標籤，優化效能
            for key, val in replacements.items():
                if key in p.text:
                    # 關鍵：不改 p.text，而是改 p.runs
                    for run in p.runs:
                        if key in run.text:
                            print(f"找到標籤 {key}，正在保留格式替換...")
                            run.text = run.text.replace(key, val)

    def replace_in_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)

    # 3. 執行替換：內文段落
    replace_in_paragraphs(doc.paragraphs)

    # 4. 執行替換：表格
    replace_in_tables(doc.tables)

    # 5. 執行替換：頁首與頁尾
    for section in doc.sections:
        replace_in_paragraphs(section.header.paragraphs)
        replace_in_tables(section.header.tables)
        replace_in_paragraphs(section.footer.paragraphs)
        replace_in_tables(section.footer.tables)

    # 6. 儲存
    doc.save(output_path)
    print(f"替換程序執行完畢，結果已儲存至：{output_path}")

# 使用範例
run_based_replace("template.docx", "data.xlsx", "run_method_result.docx")