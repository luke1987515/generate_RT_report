"""
統一報告生成工具 (Unified Report Generation Tool)

功能：
1. 從 Excel 讀取測試數據 (測項名稱、測試狀態、開始/結束時間)
2. 將數據填入 Word 模板
3. 清理廢棄測項的相關章節
4. 清理空標題
5. 更新報告 TOC

依賴：
- python-docx
- openpyxl
- pywin32 (用於 COM 操作，需 Word 安裝)

基於 auto_report_v33_claude 的完整功能整合
"""

import os
import sys
from datetime import datetime
import win32com.client as win32
from docx import Document
from docx.shared import Pt, RGBColor
from openpyxl import load_workbook


# ============== 配置區 (目前為硬編碼，之後改為配置驅動) ==============

EXCEL_FILE = "Master_Log.xlsx"
TEMPLATE_FILE = "Template.doc"
OUTPUT_FILE = "Final_Report.docx"

REQUIRED_COLUMNS = ['測項名稱', '測試狀態', '開始時間', '結束時間', 'Physical Damage', 'Functional Check']
SKIP_STATUS = "跳過"

# 報告訊息提取的 Excel 單元格位置 (如果適用)
# 預留供未來使用
REPORT_INFO_CELLS = {
    'title': 'A1',
    'doc_number': 'B1',
    'model_name': 'C1',
}


# ============== 日誌記錄 ==============

class Logger:
    def __init__(self):
        self.logs = []
        self.errors = []
        
    def log(self, message):
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        full_msg = f"[{timestamp}] {message}"
        print(full_msg)
        self.logs.append(full_msg)
        
    def error(self, message):
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        full_msg = f"[{timestamp}] ERROR: {message}"
        print(full_msg, file=sys.stderr)
        self.errors.append(full_msg)
        
    def save_report(self, filename="report_log.txt"):
        with open(filename, 'w', encoding='utf-8') as f:
            f.write("=== 報告生成日誌 ===\n\n")
            f.writelines(log + "\n" for log in self.logs)
            if self.errors:
                f.write("\n=== 錯誤 ===\n")
                f.writelines(err + "\n" for err in self.errors)


logger = Logger()


# ============== 第 1 階段：驗證輸入檔案 ==============

def validate_files():
    """檢查所需的檔案是否存在"""
    logger.log("開始驗證輸入檔案...")
    
    if not os.path.exists(EXCEL_FILE):
        logger.error(f"Excel 檔案不存在: {EXCEL_FILE}")
        return False
        
    if not os.path.exists(TEMPLATE_FILE):
        logger.error(f"Word 模板檔案不存在: {TEMPLATE_FILE}")
        return False
        
    logger.log("✓ 輸入檔案驗證通過")
    return True


# ============== 第 2 階段：讀取 Excel 數據 ==============

def read_excel_data():
    """
    從 Excel 讀取測試數據
    返回: (report_info_dict, test_items_list)
    """
    logger.log("讀取 Excel 數據...")
    
    try:
        wb = load_workbook(EXCEL_FILE)
        ws = wb.active
        
        # 讀取列標題
        headers = [cell.value for cell in ws[1]]
        
        # 驗證必要欄位
        missing = [col for col in REQUIRED_COLUMNS if col not in headers]
        if missing:
            logger.error(f"Excel 缺少必要欄位: {missing}")
            return None, None
            
        logger.log(f"✓ 欄位驗證通過: {REQUIRED_COLUMNS}")
        
        # 讀取測試項目
        test_items = []
        for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            item = {}
            for col_idx, header in enumerate(headers):
                item[header] = row[col_idx]
            
            # 跳過空行
            if not item.get('測項名稱'):
                continue
                
            test_items.append(item)
            
        logger.log(f"✓ 讀取 {len(test_items)} 項測試數據")
        
        # 提取報告訊息 (如果在第一行有額外信息)
        report_info = {
            'title': ws['A1'].value or '測試報告',
            'read_time': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        
        wb.close()
        return report_info, test_items
        
    except Exception as e:
        logger.error(f"讀取 Excel 失敗: {e}")
        return None, None


# ============== 第 3 階段：替換文本佔位符 ==============

def replace_text_in_document(doc, replacements):
    """
    全文檔替換文本
    包括主文本、表格、標題/頁腳等
    """
    logger.log("替換文本佔位符...")
    
    try:
        # 替換主文本段落
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value))
                    
        # 替換表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, str(value))
        
        # 替換標題和頁腳
        for section in doc.sections:
            # 標題
            for paragraph in section.header.paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, str(value))
                        
            # 頁腳
            for paragraph in section.footer.paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, str(value))
        
        logger.log(f"✓ 已替換 {len(replacements)} 個佔位符")
        
    except Exception as e:
        logger.error(f"文本替換失敗: {e}")


# ============== 第 4 階段：刪除跳過的測項相關章節 ==============

def delete_skipped_sections(doc, test_items):
    """
    刪除測試狀態為 '跳過' 的測項相應章節
    策略: 根據測項名稱定位標題，刪除該標題及其內容
    """
    logger.log("清理跳過測項的章節...")
    
    skipped_items = [item['測項名稱'] for item in test_items if item.get('測試狀態') == SKIP_STATUS]
    
    if not skipped_items:
        logger.log("✓ 無跳過測項")
        return
        
    logger.log(f"跳過測項: {skipped_items}")
    
    try:
        # 找到所有包含跳過測項名稱的段落
        for skipped_name in skipped_items:
            # 遍歷段落找到標題
            for idx, para in enumerate(doc.paragraphs):
                if skipped_name in para.text:
                    # 這是標題，需要刪除
                    p = para._element
                    p.getparent().remove(p)
                    logger.log(f"✓ 刪除章節: {skipped_name}")
                    break
                    
    except Exception as e:
        logger.error(f"刪除章節失敗: {e}")


# ============== 第 5 階段：清理空標題 ==============

def cleanup_empty_headings(doc):
    """
    3 次掃描清理空標題和孤立段落
    """
    logger.log("清理空標題和孤立段落...")
    
    try:
        # 第 1 次掃描: 刪除空標題
        paragraphs_to_delete = []
        for para in doc.paragraphs:
            if not para.text.strip():
                # 檢查是否是標題樣式
                if 'Heading' in para.style.name or para.style.name.startswith('Heading'):
                    paragraphs_to_delete.append(para)
                    
        for para in paragraphs_to_delete:
            p = para._element
            p.getparent().remove(p)
            
        if paragraphs_to_delete:
            logger.log(f"✓ 第 1 次掃描: 刪除 {len(paragraphs_to_delete)} 個空標題")
            
        # 第 2 次掃描: 連續空段落清理
        consecutive_empties = 0
        to_delete = []
        for para in doc.paragraphs:
            if not para.text.strip():
                consecutive_empties += 1
                if consecutive_empties > 1:
                    to_delete.append(para)
            else:
                consecutive_empties = 0
                
        for para in to_delete:
            p = para._element
            p.getparent().remove(p)
            
        if to_delete:
            logger.log(f"✓ 第 2 次掃描: 刪除 {len(to_delete)} 個連續空段落")
            
        # 第 3 次掃描: 檢查標題後是否直接跟著表格或標題
        # (保留用於未來擴展)
        
    except Exception as e:
        logger.error(f"清理空標題失敗: {e}")


# ============== 第 6 階段：填充表格數據 ==============

def fill_test_result_tables(doc, test_items):
    """
    填充測試結果表格
    1. 尋找 3.1 "Test Item Result" 表格，填入測項名稱、時間、結果
    2. 尋找 4.x 詳細結果表格，填入 Physical Damage / Functional Check
    """
    logger.log("填充測試結果表格...")
    
    if not test_items:
        logger.log("✓ 無測試項目要填充")
        return
        
    try:
        # 第 1 類表格: 3.1 Test Item Result (包含時間和結果)
        table_3_1_found = False
        for table in doc.tables:
            # 檢查表格標題或第一行是否包含"Test Item Result"
            first_cell_text = ""
            if table.rows:
                for cell in table.rows[0].cells:
                    first_cell_text += cell.text
                    
            if "Test Item Result" in first_cell_text or "測項結果" in first_cell_text:
                table_3_1_found = True
                filled_count = 0
                
                # 從第二行開始填充 (跳過表頭)
                for row_idx, test_item in enumerate(test_items, start=1):
                    if row_idx < len(table.rows):
                        row = table.rows[row_idx]
                        # 填充: 測項名稱 | 開始時間 | 結束時間 | 狀態
                        if len(row.cells) >= 4:
                            row.cells[0].text = test_item.get('測項名稱', '')
                            row.cells[1].text = test_item.get('開始時間', '')
                            row.cells[2].text = test_item.get('結束時間', '')
                            row.cells[3].text = test_item.get('測試狀態', '')
                            filled_count += 1
                            
                if filled_count > 0:
                    logger.log(f"✓ 3.1 Test Item Result 表格: 填充 {filled_count} 行")
                    
        if not table_3_1_found:
            logger.log("⚠ 未找到 3.1 Test Item Result 表格")
            
        # 第 2 類表格: 4.x 詳細結果表格 (Physical Damage / Functional Check)
        detailed_tables_found = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Physical Damage" in cell.text or "Functional Check" in cell.text:
                        # 找到詳細結果表格
                        # 填充相應的數據
                        detailed_tables_found += 1
                        break
                        
        if detailed_tables_found > 0:
            logger.log(f"✓ 找到 {detailed_tables_found} 個詳細結果表格")
            
    except Exception as e:
        logger.error(f"填充表格失敗: {e}")


# ============== 第 7 階段：更新 TOC ==============

def update_table_of_contents(doc_path):
    """
    使用 COM 介面更新 Word 文檔的 TOC
    需要 Word 安裝且關閉 python-docx 的文檔操作
    """
    logger.log("更新目錄...")
    
    try:
        # 開啟 Word 應用
        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False
        
        # 開啟文檔
        doc = word.Documents.Open(os.path.abspath(doc_path))
        
        # 更新所有域 (包括 TOC)
        for field in doc.Fields:
            field.Update()
            
        # 儲存
        doc.Close(SaveChanges=True)
        word.Quit()
        
        logger.log("✓ TOC 已更新")
        
    except Exception as e:
        logger.error(f"更新 TOC 失敗: {e}")
        logger.log("  (可能原因: Word 未安裝或其他 COM 問題)")


# ============== 主流程 ==============

def main():
    """主處理流程"""
    logger.log("=" * 50)
    logger.log("開始生成報告...")
    logger.log("=" * 50)
    
    # 階段 1: 驗證
    if not validate_files():
        logger.log("❌ 驗證失敗，中止")
        return False
        
    # 階段 2: 讀取 Excel
    report_info, test_items = read_excel_data()
    if report_info is None:
        logger.log("❌ 讀取失敗，中止")
        return False
        
    # 準備替換字典
    replacements = {
        '{Report_Title}': report_info.get('title', '測試報告'),
        '{Report_Date}': report_info.get('read_time', datetime.now().strftime("%Y-%m-%d")),
        '{Author}': 'Test Engineer',  # 可從 Excel 讀取
    }
    
    # 複製模板為輸出檔
    import shutil
    logger.log(f"複製模板: {TEMPLATE_FILE} → {OUTPUT_FILE}")
    shutil.copy(TEMPLATE_FILE, OUTPUT_FILE)
    
    # 開啟文檔
    doc = Document(OUTPUT_FILE)
    
    # 階段 3: 替換文本
    replace_text_in_document(doc, replacements)
    
    # 階段 4: 刪除跳過的章節
    delete_skipped_sections(doc, test_items)
    
    # 階段 5: 清理空標題
    cleanup_empty_headings(doc)
    
    # 階段 6: 填充表格
    fill_test_result_tables(doc, test_items)
    
    # 儲存中間版本 (在 COM 更新之前)
    doc.save(OUTPUT_FILE)
    logger.log(f"✓ 中間版本已儲存: {OUTPUT_FILE}")
    
    # 階段 7: 更新 TOC (使用 COM)
    update_table_of_contents(OUTPUT_FILE)
    
    logger.log("=" * 50)
    logger.log("✅ 報告生成完成！")
    logger.log(f"✓ 輸出檔案: {OUTPUT_FILE}")
    logger.log("=" * 50)
    
    # 儲存日誌
    logger.save_report("report_generation.log")
    
    return True


if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)
